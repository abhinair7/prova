"""
Data Loaders — ingest PDFs, CSVs, DOCX, plain text, and HTML
into a unified Document schema.
"""

import io
import csv
import logging
from pathlib import Path
from typing import Union
from app.schemas.models import Document

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Unified loader that auto-detects format and returns Document objects."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".md", ".html"}

    def load_bytes(self, content: bytes, filename: str) -> list[Document]:
        ext = Path(filename).suffix.lower()
        if ext == ".pdf":
            return self._load_pdf(content, filename)
        elif ext == ".docx":
            return self._load_docx(content, filename)
        elif ext == ".csv":
            return self._load_csv(content, filename)
        elif ext in (".txt", ".md"):
            return self._load_text(content, filename)
        elif ext == ".html":
            return self._load_html(content, filename)
        else:
            logger.warning("Unsupported extension %s — loading as plain text", ext)
            return self._load_text(content, filename)

    def load_text(self, text: str, source: str = "inline") -> list[Document]:
        return [Document(content=text.strip(), source=source, metadata={"type": "text"})]

    # ── PDF ─────────────────────────────────────────────────────────────────────
    def _load_pdf(self, content: bytes, filename: str) -> list[Document]:
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            docs = []
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    docs.append(Document(
                        content=text.strip(),
                        source=filename,
                        page=page_num + 1,
                        metadata={"type": "pdf", "page": page_num + 1, "total_pages": len(reader.pages)},
                    ))
            logger.info("Loaded %d pages from PDF: %s", len(docs), filename)
            return docs
        except ImportError:
            logger.error("pypdf not installed — pip install pypdf")
            raise

    # ── DOCX ────────────────────────────────────────────────────────────────────
    def _load_docx(self, content: bytes, filename: str) -> list[Document]:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            full_text = "\n\n".join(paragraphs)
            return [Document(
                content=full_text,
                source=filename,
                metadata={"type": "docx", "paragraph_count": len(paragraphs)},
            )]
        except ImportError:
            logger.error("python-docx not installed — pip install python-docx")
            raise

    # ── CSV ─────────────────────────────────────────────────────────────────────
    def _load_csv(self, content: bytes, filename: str) -> list[Document]:
        text = content.decode("utf-8", errors="replace")
        reader = csv.DictReader(io.StringIO(text))
        rows = list(reader)
        docs = []
        for i, row in enumerate(rows):
            row_text = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
            docs.append(Document(
                content=row_text,
                source=filename,
                metadata={"type": "csv", "row": i, "columns": list(row.keys())},
            ))
        logger.info("Loaded %d rows from CSV: %s", len(docs), filename)
        return docs

    # ── Plain text / Markdown ────────────────────────────────────────────────────
    def _load_text(self, content: bytes, filename: str) -> list[Document]:
        text = content.decode("utf-8", errors="replace").strip()
        return [Document(
            content=text,
            source=filename,
            metadata={"type": "text", "char_count": len(text)},
        )]

    # ── HTML ────────────────────────────────────────────────────────────────────
    def _load_html(self, content: bytes, filename: str) -> list[Document]:
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content, "lxml")
            for tag in soup(["script", "style", "nav", "footer"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
            return [Document(
                content=text,
                source=filename,
                metadata={"type": "html"},
            )]
        except ImportError:
            return self._load_text(content, filename)
